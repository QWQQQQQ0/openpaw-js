"""Word document generator using python-docx."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass
class TextSegment:
    text: str
    bold: bool = False
    italic: bool = False


@dataclass
class DocumentSection:
    heading: str | None = None
    heading_level: int = 2
    paragraphs: list[str | list[TextSegment]] = field(default_factory=list)


class WordGenerator:
    """Generate Word documents from structured content or Markdown."""

    def generate(
        self,
        title: str,
        content: str,
        subtitle: str | None = None,
        author: str | None = None,
    ) -> bytes:
        """Generate a Word document from Markdown content.

        Args:
            title: Document title
            content: Markdown content
            subtitle: Optional subtitle
            author: Optional author name

        Returns:
            Document as bytes
        """
        sections = self.parse_markdown(content)
        return self._build_docx(title, sections, subtitle, author)

    def generate_from_sections(
        self,
        title: str,
        sections: list[dict],
        subtitle: str | None = None,
        author: str | None = None,
    ) -> bytes:
        """Generate from structured sections.

        Args:
            title: Document title
            sections: List of {heading?, heading_level?, paragraphs: [str]}
            subtitle: Optional subtitle
            author: Optional author name

        Returns:
            Document as bytes
        """
        parsed = []
        for s in sections:
            parsed.append(DocumentSection(
                heading=s.get("heading"),
                heading_level=s.get("heading_level", 2),
                paragraphs=s.get("paragraphs", []),
            ))
        return self._build_docx(title, parsed, subtitle, author)

    def parse_markdown(self, content: str) -> list[DocumentSection]:
        """Parse Markdown content into document sections.

        Supports:
            - # / ## / ### headings
            - Plain paragraphs
            - **bold** and *italic* inline formatting
            - Unordered lists (- item)
            - Ordered lists (1. item)

        Args:
            content: Markdown string

        Returns:
            List of DocumentSection
        """
        sections: list[DocumentSection] = []
        current_section = DocumentSection()

        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            # Heading
            heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
            if heading_match:
                # Save previous section if it has content
                if current_section.paragraphs or current_section.heading:
                    sections.append(current_section)

                level = len(heading_match.group(1))
                current_section = DocumentSection(
                    heading=heading_match.group(2).strip(),
                    heading_level=level,
                )
                i += 1
                continue

            # Unordered list
            list_match = re.match(r"^[-*]\s+(.+)$", line)
            if list_match:
                items = [list_match.group(1).strip()]
                while i + 1 < len(lines):
                    next_match = re.match(r"^[-*]\s+(.+)$", lines[i + 1])
                    if next_match:
                        items.append(next_match.group(1).strip())
                        i += 1
                    else:
                        break
                for item in items:
                    current_section.paragraphs.append(f"• {item}")
                i += 1
                continue

            # Ordered list
            ordered_match = re.match(r"^(\d+)[.)]\s+(.+)$", line)
            if ordered_match:
                items = [(1, ordered_match.group(2).strip())]
                idx = 2
                while i + 1 < len(lines):
                    next_match = re.match(r"^(\d+)[.)]\s+(.+)$", lines[i + 1])
                    if next_match:
                        items.append((idx, next_match.group(2).strip()))
                        idx += 1
                        i += 1
                    else:
                        break
                for num, item in items:
                    current_section.paragraphs.append(f"{num}. {item}")
                i += 1
                continue

            # Empty line
            if not line.strip():
                i += 1
                continue

            # Regular paragraph
            current_section.paragraphs.append(line.strip())
            i += 1

        # Save last section
        if current_section.paragraphs or current_section.heading:
            sections.append(current_section)

        return sections

    def _build_docx(
        self,
        title: str,
        sections: list[DocumentSection],
        subtitle: str | None,
        author: str | None,
    ) -> bytes:
        """Build a .docx document from parsed sections."""
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.size = Pt(12)

        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        if subtitle:
            sub_para = doc.add_paragraph(subtitle)
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.runs[0].italic = True

        # Author
        if author:
            author_para = doc.add_paragraph(f"作者: {author}")
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add separator
        doc.add_paragraph("─" * 50)

        # Sections
        for section in sections:
            if section.heading:
                doc.add_heading(section.heading, level=min(section.heading_level, 3))

            for para in section.paragraphs:
                if isinstance(para, str):
                    self._add_formatted_paragraph(doc, para)
                elif isinstance(para, list):
                    p = doc.add_paragraph()
                    for seg in para:
                        run = p.add_run(seg.text)
                        run.bold = seg.bold
                        run.italic = seg.italic

        # Serialize to bytes
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _add_formatted_paragraph(self, doc: Document, text: str) -> None:
        """Add a paragraph with inline bold/italic formatting."""
        # Check for list items
        if text.startswith("• ") or re.match(r"^\d+\.\s", text):
            p = doc.add_paragraph(text, style="List Bullet" if text.startswith("• ") else "List Number")
            return

        p = doc.add_paragraph()

        # Parse inline formatting: **bold**, *italic*
        pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))"
        for match in re.finditer(pattern, text):
            if match.group(2):  # bold
                run = p.add_run(match.group(2))
                run.bold = True
            elif match.group(3):  # italic
                run = p.add_run(match.group(3))
                run.italic = True
            elif match.group(4):  # normal
                p.add_run(match.group(4))


# Standalone usage
if __name__ == "__main__":
    gen = WordGenerator()
    md = """
## 第一章 简介

这是一个**测试文档**，用于验证Word生成功能。

### 1.1 背景

Python是一门*优雅*的编程语言。

## 第二章 功能

- 支持Markdown解析
- 支持标题层级
- 支持列表

1. 第一步
2. 第二步
3. 第三步
"""
    docx_bytes = gen.generate("测试文档", md, author="OpenPaw")
    with open("test_output.docx", "wb") as f:
        f.write(docx_bytes)
    print("Generated test_output.docx")
